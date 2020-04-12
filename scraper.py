import os
import re
import shutil

import pdfcrowd
import requests
from bs4 import BeautifulSoup
from docx import Document
from pytz import unicode
from weasyprint import HTML, CSS
from xhtml2pdf import pisa
import pdfkit
from markdownify import markdownify as md


target_dir = "recepty"
tempFile = "temp"
parentURL = 'https://www.ceskatelevize.cz/porady/12309875102-pece-cela-zeme/12527-osobni-vyzva/?fbclid=IwAR17-H6-sn5yerMBmTwlEO9Evgb3d01QJlaNsSXqsxl42_mEEnL9ZuprKDQ'


def parseHeader(recipe_soup):
    divblock = recipe_soup.find('div', attrs={'class': 'recept obsah'})
    return divblock.find('h3').contents[0]


def parseContent(recipe_soup):
    return recipe_soup.find_all('div', attrs={'class': 'recept obsah'})[1].descendants


def getFullRecipe(recipe_soup):
    return recipe_soup.find_all('div', attrs={'class': 'recept obsah'})[1].prettify()


def formatRecipe(title, content, images):
    images_html = ""
    for image in images:
        images_html = images_html + r'<img src=".' + image + '" width=33% />'

    html = """
    <html>
    <head>
    <title>
    """ + title + """
    </title>
    </head>
    <body>
    <h2>
    """ + title + """
    </h2>
    """ + content + images_html + """
    </body>
    </html>
    """

    return html


def getImagesFromGallery(recipe_soup):
    imageurls = []
    divblock = recipe_soup.find('div', attrs={'class': 'gallery_v2-holder'})
    for a in divblock.find_all('a'):
        imageurls.append("https:" + a['href'])
    print(imageurls)
    return imageurls


def downloadImages(images, path):
    imageoutputs = []
    if not os.path.exists(path):
        os.mkdir(path)
    for i in range(len(images)):
        response = requests.get(images[i], stream=True)
        full_path = path + 'img' + str(i) + '.jpg'
        imageoutputs.append(full_path)
        with open(full_path, 'wb') as out_file:
            shutil.copyfileobj(response.raw, out_file)
        del response
    return imageoutputs


# https://python-docx.readthedocs.io/en/latest/
def createDocxFile(header, content, images):
    document = Document()
    document.add_heading(header, 0)
    # document.add_paragraph(BeautifulSoup(content, 'html.parser').get_text())
    # document.add_paragraph(content.strings.replace("<br>", "\\n"))

    for content_part in content:
        # if content_part.name == "blockquote" or content_part.name == "p":
        #     print(content_part.string)
        if content_part.name == "ul" or content_part.name == "ol":
            for bullet in content_part.find_all('li'):
                print("-" + unicode(bullet.string))
    # content_soup = BeautifulSoup(content_part,mys 'html.parser')
    # document.add_paragraph(content_soup.get_text())
    # document.add_paragraph(content_part.replace("<br>", "\\n"))


page = requests.get(parentURL)
soup = BeautifulSoup(page.content, 'html.parser')
results = soup.find_all('div', attrs={'class': 'explanatory-content'})
urls = []
for result in results:
    urls.append("https://www.ceskatelevize.cz" + result.find('a')['href'])
i = 0


def getRecipesAsHtml(i,title,recipe_soup):
    f = open("./temp/" + str(i) + ".html", "w", encoding="utf-8")
    images = getImagesFromGallery(recipe_soup)
    downloaded_images = downloadImages(images, "./temp/" + str(i) + "/")
    formatted = formatRecipe(title, getFullRecipe(recipe_soup), downloaded_images)
    f.write(formatted)
    f.close()
    return formatted

for url in urls:
    print(url)
    # HTML(url).write_pdf('./recepty/' + str(i) + ".pdf", stylesheets=[CSS(string='header{display:none} .grid-lg-push-1 { margin-left: 0%;}')])
    i = i + 1
    recipe_page = requests.get(url)
    recipe_soup = BeautifulSoup(recipe_page.content, 'html.parser')
    # recipe_content_all = recipe_soup.find_all('div', attrs={'class': 'recept obsah'})
    header = parseHeader(recipe_soup)
    # content = parseContent(recipe_soup)
    title = parseHeader(recipe_soup)
    # f = open("./temp/" + str(i) + ".html", "w", encoding="utf-8")
    # HTML("./temp/" + str(i) + ".html").write_pdf('./recepty/' + str(i) + ".pdf")
    # kwargs = {'bypass_robots': True, 'project_name': 'recognizable_name'}

    # save_webpage(url, "./recepty_full/" + title + "/", **kwargs)
    #   config.setup_config(url,"./recepty_full/" + title + "/", **kwargs)
    #   wp = WebPage()
    #   wp.get(url)
    # wp.save_html("./recepty_full/" + title + "/" + str(i) + title + ".html",True)
    # wp.save_assets()
    # wp.save_complete()
    # f.write(unicode(content.text))


    # getRecipesAsHtml(i,title,recipe_soup)


    # f.write(unicode(getFullRecipe(recipe_soup)))
    if not os.path.exists(target_dir):
        os.mkdir(target_dir)
    HTML("./temp/" + str(i) + ".html").write_pdf('./recepty/' + str(i) + ".pdf")
    pdfkit.from_file("./temp/" + str(i) + ".html", './recepty/' + str(i) + ".pdf")

    # terminate after one passing just for test

    #805899dc4e8191c1730f9353c2ef7177
    exit(0)
    # for contentFragment in parseContent(recipe_soup):
    #     # print(unicode(contentFragment.string))
    #     print(contentFragment)
    #     print("--------------------------")


    # if os.path.exists(tempFile):
    #     shutil.rmtree(tempFile)
    # os.mkdir("%s" % tempFile)
    # createDocxFile(header, content, images)
    # exit(0)
    # createDocxFile(header,c
    # createDocxFile(header,content,images)
    # for recipe_content in recipe_content_all:
    #     print(recipe_content.prettify())
